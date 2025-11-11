import matplotlib.pyplot as plt 
import docx 
# ------------------------------------------------------
doc = docx.Document('lion.docx')    #открываем документ 
text = [] 
letters = {} 
 # ------------------------------------------------------
word = str(input()) 
letter = str(input()) 
 # ------------------------------------------------------
for pr in doc.paragraphs:     
    text.append(pr.text) 
 
text = " ".join(text).lower() 
total_words = text.count(word) 
 
perc_1 = (total_words / len(text.split()))*100 
 
 
for i in text:    #цикл для нахождения букв в тексте 
    for l in i: 
        if l in letters: 
            letters[l] += 1 
         
        else: 
            letters.update({l:1}) 
            
 
keys = list(letters.keys())    #создаем оси для графика 
values = list(letters.values()) 
 
plt.bar(keys, values)     #создаем график 
 
 
plt.xlabel("Буквы")     
plt.ylabel("Количество") 
plt.title("Гистограмма количества букв") 
 
plt.show()    #показ графика 
 
table = doc.add_table(rows = 2, cols = 3)    #создаем таблицу 
 
table.cell(0, 0).text = "Слово" 
table.cell(0, 1).text = "Частота встречи, раз" 
table.cell(0, 2).text = "Частота встречи в %" 
table.cell(1, 0).text = str(word) 
table.cell(1, 1).text = str(total_words) 
table.cell(1, 2).text = str(perc_1) 
 
doc.save('lion.docx')

